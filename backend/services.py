import random
import uuid
import docx
import base64
from datetime import date, datetime
from docx.oxml.shape import CT_Picture
from sqlalchemy.orm import Session
import re
from backend.crud import create_question
from backend.models import Exam, ExamQuestion, Question, User
from backend.schemas import ExamCreate, ImporterResponse, QuestionCreate
from sqlalchemy.sql import func


def parse_docx(file_path):
    doc = docx.Document(file_path)
    info = {}
    questions = []
    warnings = []

    # Process metadata
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Subject:"):
            try:
                info['subject'] = text.split(":", 1)[1].strip()
            except IndexError:
                warnings.append("Subject format is incorrect.")
        elif text.startswith("Number of Quiz:"):
            try:
                info['num_quiz'] = int(text.split(":", 1)[1].strip())
            except (ValueError, IndexError):
                warnings.append("Number of Quiz format is incorrect.")
        elif text.startswith("Lecturer:"):
            info['lecturer'] = text.split(":", 1)[1].strip()
        elif text.startswith("Date:"):
            try:
                date_str = text.split(":", 1)[1].strip()
                info['date'] = datetime.strptime(date_str, "%d-%m-%Y").date()
            except (ValueError, IndexError):
                warnings.append("Date format is incorrect.")
    if len(info) != 4:
        warnings.append("Missing header information.")

    expected_order = ['qn', 'a.', 'b.', 'c.', 'd.', 'answer:', 'mark:', 'unit:', 'mix choices:']
    table_index = 0

    for table in doc.tables:
        question = {}
        order = []

        for row in table.rows:
            if len(row.cells) != 2:
                warnings.append(f"Invalid table format in table {table_index + 1}")
                continue

            key = row.cells[0].text.strip().lower()
            value = row.cells[1].text.strip()
            order.append(key)
            
            if value is None:
                warnings.append(f"Emty cell value was found in table {table_index + 1}")
                continue
            
            

            if key.strip().lower().startswith('qn'):
                qn = re.findall("^qn=(\d+)", key.strip())
                if qn[0] is None:
                    warnings.append(f"Question number not found in table {table_index + 1}")
                    continue
                question['question_number'] = qn[0]

                question['text'] = value
                # Handle image
                for paragraph in row.cells[1].paragraphs:
                    for run in paragraph.runs:
                        image_parts = []
                        for element in run._element.iter():
                            if isinstance(element, CT_Picture):
                                rId = element.xpath('.//a:blip/@r:embed')[0]
                                image_part = doc.part.related_parts[rId]
                                image_parts.append(image_part)
                    
                        if image_parts:
                            image_bytes = image_parts[0].blob
                            question['image'] = base64.b64encode(image_bytes).decode('utf-8')
                            # print('question image: ', question['image'])
                question['choices'] = [value]
            elif key.strip().lower() in ['b.', 'c.', 'd.']:
                if 'choices' in question:
                    question['choices'].append(value)
                else:
                    warnings.append(f"Choice {key} found before 'a.' in table {table_index + 1}")
            elif key.strip().lower() == 'answer:':
                question['correct_answer'] = value
            elif key.strip().lower() == 'mark:':
                try:
                    question['mark'] = float(value)
                except ValueError:
                    warnings.append(f"Invalid mark value in table {table_index + 1}")
            elif key.strip().lower() == 'unit:':
                question['unit'] = value
            elif key.strip().lower() == 'mix choices:':
                question['mix_choices'] = value.lower() == 'yes'

        # Ensure keys are in the expected order
        order = list(map(lambda x: x.strip().lower(), order))
        order[0] = order[0][:2]
        
        if order != expected_order:
            warnings.append(f"Row names are not following the expected order in table {table_index + 1}, {order}")

        if 'text' not in question:
            warnings.append(f"Question text is missing in table {table_index + 1}")
        if 'choices' not in question or len(question.get('choices', [])) < 2:
            warnings.append(f"Insufficient choices in table {table_index + 1}")
        if 'correct_answer' not in question:
            warnings.append(f"Correct answer is missing in table {table_index + 1}")

        if 'text' in question and 'choices' in question and 'correct_answer' in question:
            questions.append(question)
        
        table_index += 1
    return info, questions, warnings

def process_uploaded_file(user_id: uuid.UUID, file_path: str, db: Session):
    info, questions, warnings = parse_docx(file_path)
    


    if warnings:
        # If there are warnings, don't add questions to the database
        return ImporterResponse(
            status="fail",
            message="Warnings found during processing. Questions not added to database.",
            exam_subject=info.get('subject', ''),
            questions_created=0,
            warnings=warnings
        )

    questions_created = []
    for q in questions:
        db_question = create_question(db, QuestionCreate(
            question_number=q['question_number'],
            exam_subject=info['subject'],
            exam_maker=user_id,
            question_date=date.today(),
            question_content=q['text'],
            question_image=q.get('image'),
            option_a=q['choices'][0],
            option_b=q['choices'][1],
            option_c=q['choices'][2],
            option_d=q['choices'][3],
            correct_answer=q['correct_answer'],
            question_mark=int(q['mark']),
            question_unit=q.get('unit', ''),
            question_mixchoices=q.get('mix_choices', False)
        ))
        questions_created.append(db_question)
    
    return ImporterResponse(
        status="success",
        message="File processed successfully. Questions added to database.",
        exam_subject=info['subject'],
        questions_created=len(questions_created),
        warnings=[]
        )
        
def create_exam(db:Session, exam_data: ExamCreate):
    new_exam = Exam(
    exam_id=str(uuid.uuid4()),
    exam_subject=exam_data.exam_subject,
    exam_code=exam_data.exam_code,
    duration=exam_data.duration,
    number_of_questions=exam_data.number_of_questions)
    db.add(new_exam)
    db.commit()
        # Select questions for the exam
    questions = select_questions(db, exam_data.exam_subject, exam_data.number_of_questions)

    # Randomize answer order for each question
    randomized_questions = randomize_answers(questions)

    # Associate questions with the exam
    associate_questions_with_exam(db, new_exam.exam_id, randomized_questions)

    return new_exam

def select_questions(db: Session, subject: str, num_questions: int):
    questions = db.query(Question).filter_by(Question.exam_subject == subject).order_by(func.random()).limit(num_questions).all()
    return questions

def randomize_answers(questions):
    randomized_questions = []
    for question in questions:
        options = [('a', question.option_a),
                   ('b', question.option_b),
                   ('c', question.option_c),
                   ('d', question.option_d)]
        random.shuffle(open)

        new_question = question.copy()
        new_question.option_a = options[0][1]
        new_question.option_b = options[1][1]
        new_question.option_c = options[2][1]
        new_question.option_d = options[3][1]
        

        # Update the correct answer to match the new order
        index = 0
        for letter, option in options:
            if letter == question.correct_answer.lower():
                answers = ['a','b','c','d']   
                new_question.correct_answer = answers[index]
                break
            index = index + 1
        
        randomized_questions.append(new_question)
    
def associate_questions_with_exam(db: Session, exam_id: str, questions):
    for question in questions:
        exam_question = ExamQuestion(
            id=uuid.uuid4(),
            exam_id=exam_id,
            question_id=question.question_id,
            option_a=question.option_a,
            option_b=question.option_b,
            option_c=question.option_c,
            option_d=question.option_d,
            correct_answer=question.correct_answer
        )
        db.add(exam_question)
    db.commit()
