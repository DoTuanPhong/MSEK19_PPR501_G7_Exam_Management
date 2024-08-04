import base64
import os
from docx import Document
from app.models import Question, Choice, Subject, User
from sqlalchemy.orm import Session

class DocxProcessor:
    def __init__(self, db: Session):
        self.db = db

    def process_docx(self, file):
        doc = Document(file)
        subject_name = None
        num_quiz = None
        lecturer = None
        date = None
        questions = []

        # Process the first part
        for paragraph in doc.paragraphs[:4]:
            line = paragraph.text.strip()
            if line.startswith("Subject:"):
                subject_name = line.split(":")[1].strip()
            elif line.startswith("Number of Quiz:"):
                num_quiz = int(line.split(":")[1].strip())
            elif line.startswith("Lecturer:"):
                lecturer = line.split(":")[1].strip()
            elif line.startswith("Date:"):
                date = line.split(":")[1].strip()

        if not all([subject_name, num_quiz, lecturer, date]):
            raise ValueError("Missing required information in the document header")

        # Process the tables
        for table in doc.tables:
            question = self.process_question_table(table)
            if question:
                questions.append(question)

        subject = self.get_or_create_subject(subject_name, lecturer)
        self.save_questions(subject, questions)

        return {
            "subject": subject_name,
            "lecturer": lecturer,
            "date": date,
            "num_quiz": num_quiz,
            "questions_processed": len(questions)
        }

    def process_question_table(self, table):
        question_data = {}
        choices = []

        for row in table.rows:
            key = row.cells[0].text.strip().lower()
            value = row.cells[1].text.strip()

            if key.startswith("qn="):
                question_data["text"] = value
            elif key.startswith(("a.", "b.", "c.", "d.")):
                choices.append(value)
            elif key == "answer:":
                question_data["answer"] = value
            elif key == "mark:":
                question_data["mark"] = float(value)
            elif key == "unit:":
                question_data["unit"] = value
            elif key == "mix choices:":
                question_data["mix_choices"] = value.lower() == "yes"

        if all(k in question_data for k in ("text", "answer", "mark", "unit")) and choices:
            question_data["choices"] = choices
            return question_data
        return None


    def get_or_create_subject(self, name, lecturer_name):
        subject = self.db.query(Subject).filter(Subject.name == name).first()
        if not subject:
            lecturer = self.db.query(User).filter(User.username == lecturer_name).first()
            if not lecturer:
                raise ValueError(f"Lecturer {lecturer_name} not found")
            subject = Subject(name=name, lecturer_id=lecturer.id)
            self.db.add(subject)
            self.db.commit()
        return subject

    def save_questions(self, subject, questions):
        for q_data in questions:
            question = Question(
                subject_id=subject.id,
                text=q_data["text"],
                answer=q_data["answer"],
                mark=q_data["mark"],
                unit=q_data["unit"],
                mix_choices=q_data.get("mix_choices", True)
            )
            self.db.add(question)
            self.db.flush()  # Get the question ID

            for choice_text in q_data["choices"]:
                choice = Choice(
                    question_id=question.id,
                    text=choice_text,
                    is_correct=(choice_text == q_data["answer"])
                )
                self.db.add(choice)

        self.db.commit()