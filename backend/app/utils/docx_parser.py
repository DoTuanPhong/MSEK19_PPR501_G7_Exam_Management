import docx
import base64
from datetime import datetime
from docx.oxml.shape import CT_Picture


def parse_docx(file_path):
    doc = docx.Document(file_path)
    info = {}
    questions = []
    warnings = []

    # Process metadata
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Subject:"):
            try:
                info['subject'] = text.split(":", 1)[1].strip()
            except IndexError:
                warnings.append("Subject format is incorrect.")
        elif text.startswith("Number of Quiz:"):
            try:
                info['num_quiz'] = int(text.split(":", 1)[1].strip())
            except (ValueError, IndexError):
                warnings.append("Number of Quiz format is incorrect.")
        elif text.startswith("Lecturer:"):
            info['lecturer'] = text.split(":", 1)[1].strip()
        elif text.startswith("Date:"):
            try:
                date_str = text.split(":", 1)[1].strip()
                info['date'] = datetime.strptime(date_str, "%d-%m-%Y").date()
            except (ValueError, IndexError):
                warnings.append("Date format is incorrect.")
    if len(info) != 4:
        warnings.append("Missing header information.")

    expected_order = ['qn', 'a.', 'b.', 'c.', 'd.', 'answer:', 'mark:', 'unit:', 'mix choices:']
    table_index = 0

    for table in doc.tables:
        question = {}
        order = []

        for row in table.rows:
            if len(row.cells) != 2:
                warnings.append(f"Invalid table format in table {table_index + 1}")
                continue

            key = row.cells[0].text.strip().lower()
            value = row.cells[1].text.strip()
            order.append(key)

            if key.strip().lower().startswith('qn'):
                question['text'] = value
                # Handle image
                for paragraph in row.cells[1].paragraphs:
                    for run in paragraph.runs:
                        image_parts = []
                        for element in run._element.iter():
                            if isinstance(element, CT_Picture):
                                rId = element.xpath('.//a:blip/@r:embed')[0]
                                image_part = doc.part.related_parts[rId]
                                image_parts.append(image_part)
                    
                        if image_parts:
                            image_bytes = image_parts[0].blob
                            question['image'] = base64.b64encode(image_bytes).decode('utf-8')
                            # print('question image: ', question['image'][:50])
                question['choices'] = [value]
            elif key.strip().lower() in ['b.', 'c.', 'd.']:
                if 'choices' in question:
                    question['choices'].append(value)
                else:
                    warnings.append(f"Choice {key} found before 'a.' in table {table_index + 1}")
            elif key.strip().lower() == 'answer:':
                question['correct_answer'] = value
            elif key.strip().lower() == 'mark:':
                try:
                    question['mark'] = float(value)
                except ValueError:
                    warnings.append(f"Invalid mark value in table {table_index + 1}")
            elif key.strip().lower() == 'unit:':
                question['unit'] = value
            elif key.strip().lower() == 'mix choices:':
                question['mix_choices'] = value.lower() == 'yes'

        # Ensure keys are in the expected order
        order = list(map(lambda x: x.strip().lower(), order))
        order[0] = order[0][:2]
        
        if order != expected_order:
            warnings.append(f"Row names are not in the expected order in table {table_index + 1}, {order}")

        if 'text' not in question:
            warnings.append(f"Question text is missing in table {table_index + 1}")
        if 'choices' not in question or len(question.get('choices', [])) < 2:
            warnings.append(f"Insufficient choices in table {table_index + 1}")
        if 'correct_answer' not in question:
            warnings.append(f"Correct answer is missing in table {table_index + 1}")

        if 'text' in question and 'choices' in question and 'correct_answer' in question:
            questions.append(question)
        
    table_index += 1
    # for question in enumerate(questions,1):
    #     print(question[0], '-----------------------------------------------')
    #     for k,v in question[1].items():
    #         if k == 'image':
    #             print(k, v[:50])
    #         else:
    #             print(k,v)

    return info, questions, warnings

    return info, questions, warnings


